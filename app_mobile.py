import streamlit as st
from docx import Document
from docx.shared import Pt
import google.generativeai as genai
import io
import time

# Cấu hình trang web
st.set_page_config(page_title="Dịch Word Trung-Việt", layout="centered")

st.title("📱 App Dịch Word Song Ngữ v4.0")
st.markdown("---")

# 1. Khu vực nhập thông tin
api_key = st.text_input("Nhập Gemini API Key:", type="password", help="Lấy key tại Google AI Studio")
mode = st.radio("Lựa chọn hình thức dịch:", ("Dịch Song Ngữ (Trung trên - Việt dưới)", "Chỉ dịch sang Tiếng Việt (Thay thế)"))

# 2. Giao diện tải file
uploaded_file = st.file_uploader("Chọn tập tin Word (.docx) cần dịch", type="docx")

if uploaded_file is not None:
    if st.button("🚀 BẮT ĐẦU DỊCH NGAY"):
        if not api_key:
            st.error("Vui lòng nhập API Key trước khi dịch!")
        else:
            try:
                # Cấu hình AI
                genai.configure(api_key=api_key)
                model = genai.GenerativeModel('gemini-1.5-flash')
                
                # Đọc file Word
                doc = Document(uploaded_file)
                
                # Thu thập tất cả paragraph và ô trong bảng
                all_paras = []
                for p in doc.paragraphs:
                    if p.text.strip(): all_paras.append(p)
                for table in doc.tables:
                    for row in table.rows:
                        for cell in row.cells:
                            for p in cell.paragraphs:
                                if p.text.strip(): all_paras.append(p)

                total = len(all_paras)
                progress_bar = st.progress(0)
                st.info(f"Phát hiện {total} đoạn văn cần xử lý. Đang tiến hành dịch...")

                # Hàm xử lý dịch thuật tối ưu
                def translate_text(text):
                    prompt = f"Dịch đoạn sau sang tiếng Việt, chỉ trả về bản dịch, không giải thích: {text}"
                    # Thử lại tối đa 3 lần nếu gặp lỗi 429
                    for attempt in range(3):
                        try:
                            resp = model.generate_content(prompt)
                            return resp.text.strip()
                        except Exception as e:
                            if "429" in str(e) and attempt < 2:
                                time.sleep(5) # Nghỉ 5 giây rồi thử lại
                            else:
                                raise e
                    return None

                # Chạy dịch từng đoạn
                for i, para in enumerate(all_paras):
                    original_text = para.text
                    translated = translate_text(original_text)
                    
                    if translated:
                        if "Song ngữ" in mode:
                            # Lấy cỡ chữ của dòng gốc
                            size = 12
                            if para.runs and para.runs[0].font.size:
                                size = para.runs[0].font.size.pt
                            
                            para.add_run("\n") # Xuống dòng
                            new_run = para.add_run(translated)
                            new_run.font.name = 'Times New Roman'
                            new_run.font.size = Pt(size)
                            new_run.italic = True
                        else:
                            # Thay thế hoàn toàn
                            para.text = translated
                            for run in para.runs:
                                run.font.name = 'Times New Roman'
                    
                    # Cập nhật tiến độ
                    progress_bar.progress((i + 1) / total)

                # Lưu file vào bộ nhớ đệm
                output = io.BytesIO()
                doc.save(output)
                
                st.success("Chúc mừng! Đã dịch xong toàn bộ tài liệu.")
                st.download_button(
                    label="📥 TẢI FILE KẾT QUẢ VỀ ĐIỆN THOẠI",
                    data=output.getvalue(),
                    file_name=f"Dich_{uploaded_file.name}",
                    mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
                )

            except Exception as e:
                st.error(f"Đã xảy ra lỗi: {str(e)}")
                st.warning("Mẹo: Nếu lỗi 'Quota', hãy đợi 1 phút rồi bấm lại hoặc chia nhỏ file Word.")

st.markdown("---")
st.caption("Ứng dụng được thiết kế bởi Gemini - Hỗ trợ font Times New Roman & Giữ nguyên Format.")
